
import requests
from bs4 import BeautifulSoup
import pandas as pd
import streamlit as st
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt

def scrape_taux_change_beac_optimise():
    url = "https://www.beac.int/"
    headers = {'User-Agent': 'Mozilla/5.0'}
    response = requests.get(url, headers=headers, timeout=15)
    response.raise_for_status()

    soup = BeautifulSoup(response.content, 'html.parser')
    taux_change_list = soup.find('div', class_='taux_de_change_list')
    taux_elements = taux_change_list.find_all('div', class_='taux_de_change')

    data_list = []
    for element in taux_elements:
        left = element.find('div', id='left').get_text(strip=True)
        middle = element.find('div', id='middle').get_text(strip=True)
        right = element.find('div', id='right').get_text(strip=True)
        data_list.append({'PAIRE DEVISES': left, 'ACHAT': middle, 'VENTE': right})

    df = pd.DataFrame(data_list).iloc[1:]
    df['PAIRE DEVISES'] = df['PAIRE DEVISES'].str.replace('/XAF', '', regex=False)
    df['SYMBOLE'] = df['PAIRE DEVISES'].str.extract(r'([A-Z]{3})')

    mapping = {
        'EUR': {'CODE': '290', 'INTITULE': 'EURO'},
        'USD': {'CODE': '340', 'INTITULE': 'DOLLAR US'},
        'GBP': {'CODE': '260', 'INTITULE': 'LIVRE STERLING'},
        'CHF': {'CODE': '710', 'INTITULE': 'FRANC SUISSE'},
        'JPY': {'CODE': '350', 'INTITULE': 'YEN JAPONAIS'},
        'CAD': {'CODE': '330', 'INTITULE': 'DOLLAR CANADIEN'},
        'SEK': {'CODE': '310', 'INTITULE': 'COURONNE SUEDOISE'},
        'ZAR': {'CODE': '360', 'INTITULE': 'RAND SUD-AFRICAIN'},
        'MAD': {'CODE': '', 'INTITULE': 'DIRHAM MAROCAIN'},
        'SAR': {'CODE': '602', 'INTITULE': 'RIYAL SAOUDIEN'},
        'AED': {'CODE': '', 'INTITULE': 'DIRHAM EAU'},
        'CNY': {'CODE': '', 'INTITULE': 'RENMINBI'},
        'DKK': {'CODE': '240', 'INTITULE': 'COURONNE DANOISE'},
        'XOF': {'CODE': '002', 'INTITULE': 'FRANC CFA UEMOA'}
    }


    df['CODE'] = df['SYMBOLE'].map(lambda x: mapping.get(x, {}).get('CODE', ''))
    df['INTITULE'] = df['SYMBOLE'].map(lambda x: mapping.get(x, {}).get('INTITULE', ''))

    date_div = soup.find('div', class_='date_source_taux')
    date_source = date_div.get_text(strip=True) if date_div else 'Date inconnue'
    return df[['CODE', 'SYMBOLE', 'INTITULE', 'ACHAT', 'VENTE']].reset_index(drop=True), date_source, url

def export_to_pdf(df, date_source, url, file_path, logo_path=None):
    pdf = FPDF()
    pdf.add_page()
    if logo_path:
        pdf.image(logo_path, x=10, y=8, w=30)
        pdf.ln(20)
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Taux de Change - BEAC", ln=True, align='C')
    pdf.cell(200, 10, txt=f"Date de publication : {date_source}", ln=True)
    pdf.cell(200, 10, txt=f"Source : {url}", ln=True)
    pdf.ln(10)

    col_widths = [20, 25, 60, 30, 30]
    headers = ['CODE', 'SYMBOLE', 'INTITULE', 'ACHAT', 'VENTE']
    for i, h in enumerate(headers):
        pdf.cell(col_widths[i], 10, h, border=1)
    pdf.ln()

    for _, row in df.iterrows():
        for i, key in enumerate(headers):
            txt = str(row[key])[:18]
            pdf.cell(col_widths[i], 10, txt, border=1)
        pdf.ln()
    pdf.output(file_path)

def export_to_docx(df, date_source, url, file_path, logo_path=None):
    doc = Document()
    if logo_path:
        doc.add_picture(logo_path, width=Inches(1.0))
    doc.add_heading('Taux de Change - BEAC', level=1)
    doc.add_paragraph(f"Date de publication : {date_source}")
    doc.add_paragraph(f"Source : {url}")
    doc.add_paragraph(" ")

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        run = hdr_cells[i].paragraphs[0].add_run(col)
        run.bold = True

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.save(file_path)

def main():
    st.title("💱 Taux de change - BEAC")
    st.write("Application de scraping automatique des taux de change publiés par la BEAC.")
    df, date_source, source_url = scrape_taux_change_beac_optimise()
    st.success("Scraping terminé avec succès !")
    st.markdown(f"**📅 ** {date_source}")
    st.markdown(f"**🔗 Source :** [Site de la BEAC]({source_url})")
    st.dataframe(df, use_container_width=True)

    logo_path = "afriland_logo.jpg"  # À remplacer par ton image

    col1, col2 = st.columns(2)
    with col1:
        if st.button("📥 Télécharger en PDF"):
            pdf_path = "taux_beac.pdf"
            export_to_pdf(df, date_source, source_url, pdf_path, logo_path)
            with open(pdf_path, "rb") as f:
                st.download_button("Télécharger PDF", f, file_name=pdf_path)

    with col2:
        if st.button("📥 Télécharger en Word"):
            doc_path = "taux_beac.docx"
            export_to_docx(df, date_source, source_url, doc_path, logo_path)
            with open(doc_path, "rb") as f:
                st.download_button("Télécharger DOCX", f, file_name=doc_path)

if __name__ == '__main__':
    main()
